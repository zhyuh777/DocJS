#!/usr/bin/env python3
"""Smart Teaching Plan Generator — pure Python, uses python-docx + DeepSeek."""
import json, shutil, sys, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests

# ============================================================
# Config — edit these or pass via command line
# ============================================================
API_KEY = os.environ.get("DEEPSEEK_KEY", "")
API_URL = "https://api.deepseek.com/v1"
MODEL = "deepseek-chat"
TEMPLATE_PATH = None  # set by command line
DESCRIPTION = None    # set by command line


def extract_structure(doc):
    """Read template structure — paragraphs and tables with their text."""
    blocks = []
    for pi, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        level = None
        for h in range(1, 8):
            if f"Heading {h}" in style_name:
                level = h
                break
        if level:
            blocks.append({"i": len(blocks), "type": "header", "level": level, "text": text})
        else:
            blocks.append({"i": len(blocks), "type": "paragraph", "text": text})

    flat_idx = len(blocks)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                blocks.append({
                    "i": len(blocks),
                    "type": "cell",
                    "table": ti,
                    "row": ri,
                    "col": ci,
                    "text": cell.text.strip()
                })

    return blocks


def build_prompt(structure, description):
    """Build a concise system prompt for DeepSeek."""
    brief = []
    for b in structure:
        if b["type"] == "header":
            brief.append({"i": b["i"], "h": f"H{b['level']}", "t": b["text"]})
        elif b["type"] == "paragraph":
            brief.append({"i": b["i"], "p": b["text"][:100]})
        elif b["type"] == "cell" and b["row"] == 0:
            brief.append({"i": b["i"], "cell": f"t{b['table']}r{b['row']}c{b['col']}", "t": b["text"]})

    return f"""你是专业教案撰写人。根据模板结构，按用户描述填充所有内容。

模板结构（共 {len(structure)} 个位置需填充）：
{json.dumps(brief, ensure_ascii=False)}

输出 JSON：
{{"fields":[{{"index":0,"content":"填充的文本"}},{{"index":15,"content":"单元格内容"}}]}}

规则：保留标签前缀如"课题名称：xxx"。每个 index 必须填写。输出 JSON，不解释。"""


def call_deepseek(prompt, description):
    """Call DeepSeek API and return parsed fields."""
    resp = requests.post(
        f"{API_URL.rstrip('/')}/chat/completions",
        headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"},
        json={
            "model": MODEL,
            "messages": [
                {"role": "system", "content": prompt},
                {"role": "user", "content": description}
            ],
            "response_format": {"type": "json_object"},
            "max_tokens": 8192
        },
        timeout=300
    )
    resp.raise_for_status()
    data = resp.json()
    text = data["choices"][0]["message"]["content"]
    result = json.loads(text)
    return result.get("fields", [])


def fill_document(doc, fills, structure):
    """Fill content into doc by index, preserving label prefixes."""
    # Map structure index → paragraph or cell
    para_list = [p for p in doc.paragraphs if p.text.strip()]

    # Build cell flat list
    cell_list = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_list.append(cell)

    applied = 0
    for item in fills:
        idx = item.get("index", -1)
        content = item.get("content", "")
        if idx < 0 or not content:
            continue

        # Find the block
        block = None
        for b in structure:
            if b["i"] == idx:
                block = b
                break
        if not block:
            continue

        if block["type"] in ("header", "paragraph"):
            para_idx = sum(1 for b in structure[:idx] if b["type"] in ("header", "paragraph") and b["i"] < idx)
            para_idx = sum(1 for b in structure[:idx+1] if b["type"] in ("header", "paragraph")) - 1
            if 0 <= para_idx < len(para_list):
                orig = para_list[para_idx].text
                # Preserve label prefix
                if "：" in orig:
                    prefix = orig.split("：")[0] + "："
                    if not content.startswith(prefix):
                        content = prefix + content.lstrip("：")
                elif ":" in orig:
                    prefix = orig.split(":")[0] + ":"
                    if not content.startswith(prefix):
                        content = prefix + content.lstrip(":")
                para_list[para_idx].text = content
                applied += 1

        elif block["type"] == "cell":
            # Count preceding cells
            cell_idx = sum(1 for b in structure[:idx] if b["type"] == "cell")
            if 0 <= cell_idx < len(cell_list):
                cell_list[cell_idx].text = content
                applied += 1

    return applied


def main():
    global API_KEY, TEMPLATE_PATH, DESCRIPTION

    if len(sys.argv) < 2:
        print("用法: python generate.py <模板.docx> [描述]")
        print("环境变量: DEEPSEEK_KEY=sk-xxx")
        sys.exit(1)

    TEMPLATE_PATH = Path(sys.argv[1])
    if not TEMPLATE_PATH.exists():
        print(f"模板不存在: {TEMPLATE_PATH}")
        sys.exit(1)

    DESCRIPTION = sys.argv[2] if len(sys.argv) > 2 else input("请输入教案描述: ")

    # API Key
    API_KEY = os.environ.get("DEEPSEEK_KEY", "")
    if not API_KEY:
        API_KEY = input("请输入 DeepSeek API Key: ").strip()
    if not API_KEY:
        print("需要 API Key")
        sys.exit(1)

    print(f"\n📄 读取模板: {TEMPLATE_PATH.name}")
    doc = Document(str(TEMPLATE_PATH))
    structure = extract_structure(doc)
    print(f"   解析完成 — {len(structure)} 个位置（{sum(1 for b in structure if b['type']=='header')} 标题, {sum(1 for b in structure if b['type']=='paragraph')} 段落, {sum(1 for b in structure if b['type']=='cell')} 单元格）")

    print(f"\n🤖 调用 DeepSeek API ({MODEL})...")
    prompt = build_prompt(structure, DESCRIPTION)
    fills = call_deepseek(prompt, DESCRIPTION)
    print(f"   AI 返回 {len(fills)} 个字段")

    # Copy template and fill
    output_path = TEMPLATE_PATH.parent / f"{TEMPLATE_PATH.stem}_生成稿.docx"
    shutil.copy(str(TEMPLATE_PATH), str(output_path))
    doc_out = Document(str(output_path))

    print(f"\n📝 填入内容...")
    applied = fill_document(doc_out, fills, structure)
    doc_out.save(str(output_path))

    print(f"   ✓ 成功填充 {applied} 处")
    print(f"   ✓ 输出文件: {output_path}")
    print(f"\n✅ 完成！打开文件查看: {output_path}")


if __name__ == "__main__":
    main()
