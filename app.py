#!/usr/bin/env python3
"""DocJS — python-docx 直接编辑模板。
提取时同时暴露标签(已有文字)和填充位(空白格)，AI 从标签上下文理解填充内容。"""
import json, shutil, uuid, re
from pathlib import Path
from flask import Flask, request, jsonify, send_file, render_template
from docx import Document

NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
app = Flask(__name__)
UPLOAD_DIR, OUTPUT_DIR = Path(__file__).parent/"uploads", Path(__file__).parent/"outputs"
UPLOAD_DIR.mkdir(exist_ok=True); OUTPUT_DIR.mkdir(exist_ok=True)


# ============================================================
# 共享：提取表格中所有可见单元格 → (table, row, col, text, is_skip)
# ============================================================
def iter_visible_cells(doc):
    """遍历所有表格的可见单元格（跳过合并延续格）。返回生成器。"""
    for ti, t in enumerate(doc.tables):
        for ri, r in enumerate(t.rows):
            for ci, c in enumerate(r.cells):
                tcPr = c._tc.find(f'{NS}tcPr')
                if tcPr is not None:
                    hM = tcPr.find(f'{NS}hMerge')
                    vM = tcPr.find(f'{NS}vMerge')
                    # 跳过 hMerge != restart 的延续格
                    if hM is not None and hM.get(f'{NS}val') != 'restart':
                        continue
                    # 跳过 vMerge != restart 的延续格
                    if vM is not None and vM.get(f'{NS}val') != 'restart':
                        continue
                yield ti, ri, ci, c.text.strip()


def extract_structure(doc):
    """
    返回 list[dict]，完整暴露表格布局。
    规则：
      - 整行所有格文字完全相同 → 跳过（section header，如"教学过程"）
      - 非空格 → type="label"（给 AI 当上下文）
      - 空格   → type="fill"（需填内容），标签优先取同行左边最近非空格，其次取上一行列头
    """
    from collections import defaultdict

    # Step 1: 分组
    rows_data = defaultdict(list)  # (ti, ri) -> [(ci, text)]
    for ti, ri, ci, txt in iter_visible_cells(doc):
        rows_data[(ti, ri)].append((ci, txt))

    # Step 2: 收集"列头行"
    # 条件：全非空、文字不全相同、平均文本长度短（列头字少，数据行字多）
    col_headers = {}
    for (ti, ri), cells in rows_data.items():
        texts = [t for _, t in cells]
        avg_len = sum(len(t) for t in texts) / len(texts) if texts else 0
        if all(t for t in texts) and len(set(texts)) > 1 and avg_len < 6:
            col_headers[(ti, ri)] = {ci: t for ci, t in cells}

    # Step 3: 遍历每行，生成结构
    # 列头跨行持久：新 col_header 替换旧的，全相同行清空
    result = []
    current_col_headers = {}  # 当前生效的列头 {ci: text}

    for (ti, ri), cells in sorted(rows_data.items()):
        texts = [t for _, t in cells]
        nonempty = [t for t in texts if t]

        # 整行相同文字 → 清空列头（section break）
        if len(texts) > 1 and len(set(texts)) == 1:
            current_col_headers = {}
            continue
        if not nonempty:
            continue

        # 遇到新列头行 → 替换
        if (ti, ri) in col_headers:
            current_col_headers = col_headers[(ti, ri)]
        # 教学反思行 → 新段落，清空列头
        if texts and "教学反思" in texts[0]:
            current_col_headers = {}

        for ci, txt in sorted(cells, key=lambda x: x[0]):
            if txt:
                result.append({
                    "table": ti, "row": ri, "col": ci,
                    "type": "label", "label": txt[:120]
                })
            else:
                col_label = current_col_headers.get(ci, "")
                row_label = ""
                for cj, ctxt in cells:
                    if cj < ci and ctxt:
                        row_label = ctxt[:80]
                if col_label:
                    context = col_label
                elif row_label:
                    context = row_label
                else:
                    context = ""
                label = f"{context} — 需填入值" if context else "需填入值"
                result.append({
                    "table": ti, "row": ri, "col": ci,
                    "type": "fill", "label": label
                })

    # Step 4: 同行同上下文的连续空格合并为1个（避免 AI 写多份相同内容）
    # 写入时会自动展开
    merged = []
    i = 0
    while i < len(result):
        s = result[i]
        if s["type"] != "fill":
            merged.append(s)
            i += 1
            continue
        # 找到同行同上下文的连续 fill 格
        cells_to_merge = [s]
        j = i + 1
        while j < len(result) and result[j]["type"] == "fill" \
              and result[j]["table"] == s["table"] \
              and result[j]["row"] == s["row"] \
              and result[j]["label"] == s["label"]:
            cells_to_merge.append(result[j])
            j += 1
        # 合并：记录第一个格的位置 + 合并格列表
        s["merge_cells"] = [(c["table"], c["row"], c["col"]) for c in cells_to_merge]
        merged.append(s)
        i = j

    return merged


@app.route("/api/builtin-template")
def builtin_template():
    """返回内置教案模板。启动时已确保存在。"""
    tpl = UPLOAD_DIR / "builtin.docx"
    if not tpl.exists():
        return jsonify({"error":"内置模板未找到"}), 404
    doc = Document(str(tpl))
    structure = extract_structure(doc)
    block_count = sum(1 for s in structure if s["type"] == "fill")
    return jsonify({
        "doc_id": "builtin",
        "filename": "教案模板（内置）",
        "structure": structure,
        "block_count": block_count
    })


# ============================================================
# Route: 首页
# ============================================================
@app.route("/")
def index():
    return render_template("index.html")


# ============================================================
# Route: 上传模板 → 返回结构和所有可填位置
# ============================================================
@app.route("/api/upload", methods=["POST"])
def upload():
    file = request.files.get("file")
    if not file or not file.filename.endswith(".docx"):
        return jsonify({"error":"请上传.docx"}), 400
    did = uuid.uuid4().hex[:12]
    file.save(UPLOAD_DIR / f"{did}.docx")
    doc = Document(str(UPLOAD_DIR / f"{did}.docx"))

    structure = extract_structure(doc)
    block_count = sum(1 for s in structure if s["type"] == "fill")

    return jsonify({
        "doc_id": did,
        "filename": file.filename,
        "structure": structure,
        "block_count": block_count
    })


# ============================================================
# Route: AI 生成 → 填入模板 → 返回下载链接
# ============================================================
@app.route("/api/generate", methods=["POST"])
def generate():
    d = request.get_json()
    did, key = d.get("doc_id"), d.get("api_key")
    url   = d.get("api_url", "https://api.deepseek.com/v1")
    model = d.get("model", "deepseek-chat")
    desc  = d.get("description", "")
    tpl   = UPLOAD_DIR / f"{did}.docx"

    if not did or not tpl.exists():  return jsonify({"error":"模板不存在"}), 400
    if not key:                      return jsonify({"error":"请配置 API Key"}), 400
    if not desc:                     return jsonify({"error":"请输入描述"}), 400

    doc = Document(str(tpl))
    structure = extract_structure(doc)
    fills = [s for s in structure if s["type"] == "fill"]

    if not fills:
        return jsonify({"error":"模板中没有检测到可填充位置"}), 400

    # ---- 构建 AI prompt ----
    # 把所有位置（标签 + 填充位）按顺序列出，AI 从中理解上下文
    lines = []
    for idx, s in enumerate(structure):
        tag = "[标签]" if s["type"] == "label" else "→ [填充]"
        lines.append(f"{idx}. {tag} R{s['row']}C{s['col']}: {s['label']}")

    prompt = f"""你是专业教案撰写人。以下是教案模板的单元格结构（{len(structure)} 个位置，其中 {len(fills)} 个需填充）。

{chr(10).join(lines)}

请根据用户描述，为每个「→ [填充]」位置生成合适的教学内容。
输出一个 JSON 数组，只包含填充位的内容（按填充位出现的顺序），不要包含标签位：
["第1个填充位的内容","第2个填充位的内容",...]

要求：
- 高职院校教学语言，通俗易懂，短句为主
- 教学目标用布鲁姆动词（说出、列举、识别、运用、操作、调试等）
- 学情分析按"已掌握…/对…理解不深/通过…帮助…"句式
- 教学过程按"讲授→演示→实操→讨论→总结"五环节
- 每个值是纯文本，不要用 markdown
只输出 JSON 数组。"""

    import requests
    resp = requests.post(
        f"{url.rstrip('/')}/chat/completions",
        headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
        json={
            "model": model,
            "messages": [
                {"role": "system", "content": prompt},
                {"role": "user",   "content": desc}
            ],
            "max_tokens": 8192,
            "temperature": 0.7
        },
        timeout=300
    )
    if not resp.ok:
        return jsonify({"error": f"API {resp.status_code}"}), 500

    ai_text = resp.json()["choices"][0]["message"]["content"]

    # ---- 解析 AI 返回的 JSON 数组 ----
    contents = None
    try:
        contents = json.loads(ai_text)
    except:
        # 尝试提取最后一个 [...] 块
        matches = list(re.finditer(r'\[[\s\S]*\]', ai_text))
        if matches:
            try:
                contents = json.loads(matches[-1].group(0))
            except:
                pass

    if not contents or not isinstance(contents, list):
        return jsonify({"error": "AI 返回无法解析为数组", "raw": ai_text[:500]}), 500

    # ---- 填入文档 ----
    oid = uuid.uuid4().hex[:12]
    opath = OUTPUT_DIR / f"{oid}.docx"
    shutil.copy(str(tpl), str(opath))
    doc_out = Document(str(opath))
    filled = 0

    for i, s in enumerate(fills):
        if i >= len(contents) or not contents[i]:
            continue
        content = str(contents[i]).strip()
        if not content or content == "None":
            continue
        # 写入到所有合并格中
        targets = s.get("merge_cells", [(s["table"], s["row"], s["col"])])
        for ti, ri, ci in targets:
            try:
                cell = doc_out.tables[ti].rows[ri].cells[ci]
                cell.text = ""
                cell.paragraphs[0].add_run(content)
                filled += 1
            except:
                pass

    doc_out.save(str(opath))

    return jsonify({
        "output_id":     oid,
        "filled_count":  filled,
        "total_blocks":  len(fills),
        "download_url": f"/api/download/{oid}"
    })


# ============================================================
# Route: 下载生成的 docx
# ============================================================
@app.route("/api/download/<oid>")
def download(oid):
    fp = OUTPUT_DIR / f"{oid}.docx"
    if fp.exists():
        return send_file(str(fp), as_attachment=True, download_name="教案_生成版.docx")
    return jsonify({"error": "文件不存在"}), 404


if __name__ == "__main__":
    app.run(debug=True, port=5757)
